"""
Document Service for Justice-Bot
File processing and text extraction
"""

import os
import docx
import PyPDF2
from typing import Optional

def extract_text_from_file(file_path: str, content_type: str) -> Optional[str]:
    """
    Extract text content from uploaded files
    Supports: .txt, .docx, .pdf
    """
    try:
        if not os.path.exists(file_path):
            return None
        
        # Text files
        if content_type == 'text/plain' or file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        
        # Word documents
        elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                              'application/msword'] or file_path.endswith(('.docx', '.doc')):
            return extract_text_from_docx(file_path)
        
        # PDF files
        elif content_type == 'application/pdf' or file_path.endswith('.pdf'):
            return extract_text_from_pdf(file_path)
        
        else:
            return None
            
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return None

def extract_text_from_docx(file_path: str) -> Optional[str]:
    """Extract text from DOCX files"""
    try:
        doc = docx.Document(file_path)
        text_content = []
        
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        return '\n'.join(text_content)
        
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {e}")
        return None

def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """Extract text from PDF files"""
    try:
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(page_text.strip())
        
        return '\n'.join(text_content)
        
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {e}")
        return None

def generate_legal_document(template_path: str, user_data: dict, case_data: dict, output_path: str) -> bool:
    """
    Generate legal document from template
    Replace placeholders with actual data
    """
    try:
        # Load template
        doc = docx.Document(template_path)
        
        # Prepare replacement data
        replacements = {
            '{{ full_name }}': user_data.get('full_name', ''),
            '{{ email }}': user_data.get('email', ''),
            '{{ phone }}': user_data.get('phone', ''),
            '{{ address }}': user_data.get('address', ''),
            '{{ postal_code }}': user_data.get('postal_code', ''),
            '{{ province }}': user_data.get('province', ''),
            '{{ case_title }}': case_data.get('title', ''),
            '{{ issue_description }}': case_data.get('description', ''),
            '{{ legal_issue_type }}': case_data.get('legal_issue_type', ''),
            '{{ merit_score }}': str(case_data.get('merit_score', '')),
            '{{ classification }}': case_data.get('classification', ''),
            '{{ date }}': case_data.get('created_at', '').strftime('%B %d, %Y') if case_data.get('created_at') else ''
        }
        
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Save document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error generating document: {e}")
        return False

def create_default_template(template_path: str, template_type: str) -> bool:
    """
    Create default legal document templates
    """
    try:
        doc = docx.Document()
        
        if template_type == 'ltb_t2':
            # LTB T2 Application template
            doc.add_heading('Landlord and Tenant Board', 0)
            doc.add_heading('Application About a Tenant\'s Rights (Form T2)', level=1)
            
            doc.add_paragraph('Applicant Information:')
            doc.add_paragraph('Full Name: {{ full_name }}')
            doc.add_paragraph('Email: {{ email }}')
            doc.add_paragraph('Phone: {{ phone }}')
            doc.add_paragraph('Address: {{ address }}')
            doc.add_paragraph('Postal Code: {{ postal_code }}')
            doc.add_paragraph('Province: {{ province }}')
            
            doc.add_paragraph('\nCase Information:')
            doc.add_paragraph('Case Title: {{ case_title }}')
            doc.add_paragraph('Legal Issue Type: {{ legal_issue_type }}')
            doc.add_paragraph('Classification: {{ classification }}')
            doc.add_paragraph('Merit Score: {{ merit_score }}/100')
            
            doc.add_paragraph('\nIssue Description:')
            doc.add_paragraph('{{ issue_description }}')
            
            doc.add_paragraph('\nDate: {{ date }}')
            
        elif template_type == 'small_claims':
            # Small Claims Court template
            doc.add_heading('Small Claims Court', 0)
            doc.add_heading('Plaintiff\'s Claim (Form 7A)', level=1)
            
            doc.add_paragraph('Plaintiff Information:')
            doc.add_paragraph('Full Name: {{ full_name }}')
            doc.add_paragraph('Email: {{ email }}')
            doc.add_paragraph('Phone: {{ phone }}')
            doc.add_paragraph('Address: {{ address }}')
            doc.add_paragraph('Postal Code: {{ postal_code }}')
            doc.add_paragraph('Province: {{ province }}')
            
            doc.add_paragraph('\nClaim Details:')
            doc.add_paragraph('Case Title: {{ case_title }}')
            doc.add_paragraph('Legal Issue Type: {{ legal_issue_type }}')
            doc.add_paragraph('Classification: {{ classification }}')
            
            doc.add_paragraph('\nReason for Claim:')
            doc.add_paragraph('{{ issue_description }}')
            
            doc.add_paragraph('\nDate: {{ date }}')
            
        else:
            # Generic legal document template
            doc.add_heading('Legal Document', 0)
            doc.add_heading('{{ case_title }}', level=1)
            
            doc.add_paragraph('Prepared for: {{ full_name }}')
            doc.add_paragraph('Contact: {{ email }} | {{ phone }}')
            doc.add_paragraph('Address: {{ address }}, {{ postal_code }}, {{ province }}')
            
            doc.add_paragraph('\nDocument Type: {{ legal_issue_type }}')
            doc.add_paragraph('Classification: {{ classification }}')
            doc.add_paragraph('Assessment Score: {{ merit_score }}/100')
            
            doc.add_paragraph('\nDetails:')
            doc.add_paragraph('{{ issue_description }}')
            
            doc.add_paragraph('\nPrepared on: {{ date }}')
        
        # Ensure directory exists
        os.makedirs(os.path.dirname(template_path), exist_ok=True)
        doc.save(template_path)
        return True
        
    except Exception as e:
        print(f"Error creating template: {e}")
        return False