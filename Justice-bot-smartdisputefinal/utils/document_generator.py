import os
import io
import uuid
import logging
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

# Configure logging
logger = logging.getLogger(__name__)

def generate_legal_document(case, documents, form_type, form_data):
    """
    Generate a legal document based on case, documents, form type and form data
    
    Args:
        case: Case model instance
        documents: List of Document model instances
        form_type: Type of form to generate
        form_data: Dictionary of form data
        
    Returns:
        tuple: (file_path, citations)
            - file_path (str): Path to the generated document
            - citations (list): List of citations used in the document
    """
    try:
        # Determine the document generator to use based on form type
        if 'landlord-tenant' in form_type:
            return generate_ltb_form(case, documents, form_type, form_data)
        elif 'credit' in form_type:
            return generate_credit_dispute(case, documents, form_type, form_data)
        elif 'human-rights' in form_type:
            return generate_human_rights_complaint(case, documents, form_type, form_data)
        elif 'small-claims' in form_type:
            return generate_small_claims_form(case, documents, form_type, form_data)
        elif 'child-protection' in form_type:
            return generate_child_protection_form(case, documents, form_type, form_data)
        elif 'police-misconduct' in form_type:
            return generate_police_complaint(case, documents, form_type, form_data)
        elif 'referral' in form_type:
            return generate_referral_letter(case, documents, form_type, form_data)
        else:
            return generate_generic_document(case, documents, form_type, form_data)
    except Exception as e:
        logger.error(f"Error generating document: {str(e)}")
        raise

def generate_ltb_form(case, documents, form_type, form_data):
    """
    Generate a Landlord and Tenant Board form
    
    Args:
        case: Case model instance
        documents: List of Document model instances
        form_type: Type of form to generate
        form_data: Dictionary of form data
        
    Returns:
        tuple: (file_path, citations)
    """
    try:
        # Create a DocX document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = f"LTB Form - {form_type.split('_')[-1].upper()}"
        doc.core_properties.author = form_data.get('full_name', '')
        
        # Add header with LTB logo and title
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run('LANDLORD AND TENANT BOARD')
        header_run.bold = True
        header_run.font.size = Pt(16)
        
        # Add title
        if 'maintenance_issues' in form_type:
            title = "TENANT APPLICATION ABOUT MAINTENANCE (Form T6)"
            description = "Application under Section 29 of the Residential Tenancies Act, 2006"
        elif 'eviction_defense' in form_type:
            title = "TENANT APPLICATION ABOUT AN EVICTION (Form T5)"
            description = "Application under Section 82 of the Residential Tenancies Act, 2006"
        elif 'illegal_rent_increase' in form_type:
            title = "TENANT APPLICATION ABOUT A RENT INCREASE (Form T1)"
            description = "Application under Section 115 of the Residential Tenancies Act, 2006"
        elif 'harassment' in form_type:
            title = "TENANT APPLICATION ABOUT LANDLORD HARASSMENT (Form T2)"
            description = "Application under Section 29 of the Residential Tenancies Act, 2006"
        else:
            title = "TENANT APPLICATION"
            description = "Application under the Residential Tenancies Act, 2006"
        
        doc.add_heading(title, 0)
        para = doc.add_paragraph(description)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add part 1: General Information
        doc.add_heading('Part 1: General Information', 1)
        doc.add_heading('Tenant Information', 2)
        
        tenant_table = doc.add_table(rows=4, cols=2)
        tenant_table.style = 'Table Grid'
        
        # Name row
        tenant_table.cell(0, 0).text = 'Tenant Name:'
        tenant_table.cell(0, 1).text = form_data.get('full_name', '')
        
        # Address row
        tenant_table.cell(1, 0).text = 'Tenant Address:'
        tenant_table.cell(1, 1).text = f"{form_data.get('address', '')}, {form_data.get('city', '')}, {form_data.get('province', '')}, {form_data.get('postal_code', '')}"
        
        # Contact info row
        tenant_table.cell(2, 0).text = 'Phone Number:'
        tenant_table.cell(2, 1).text = form_data.get('phone', 'N/A')
        
        # Email row
        tenant_table.cell(3, 0).text = 'Email:'
        tenant_table.cell(3, 1).text = form_data.get('email', 'N/A')
        
        doc.add_paragraph()
        
        # Landlord information
        doc.add_heading('Landlord Information', 2)
        
        landlord_table = doc.add_table(rows=2, cols=2)
        landlord_table.style = 'Table Grid'
        
        # Name row
        landlord_table.cell(0, 0).text = 'Landlord Name:'
        landlord_table.cell(0, 1).text = form_data.get('landlord_name', '')
        
        # Address row
        landlord_table.cell(1, 0).text = 'Landlord Address:'
        landlord_table.cell(1, 1).text = form_data.get('landlord_address', '')
        
        doc.add_paragraph()
        
        # Rental unit information
        doc.add_heading('Rental Unit Information', 2)
        
        rental_table = doc.add_table(rows=2, cols=2)
        rental_table.style = 'Table Grid'
        
        # Address row
        rental_table.cell(0, 0).text = 'Rental Unit Address:'
        rental_table.cell(0, 1).text = form_data.get('rental_address', '')
        
        # Lease date row
        rental_table.cell(1, 0).text = 'Lease Start Date:'
        lease_start = form_data.get('lease_start', '')
        if lease_start:
            try:
                lease_start_date = datetime.strptime(lease_start, '%Y-%m-%d')
                formatted_date = lease_start_date.strftime('%B %d, %Y')
                rental_table.cell(1, 1).text = formatted_date
            except ValueError:
                rental_table.cell(1, 1).text = lease_start
        
        doc.add_paragraph()
        
        # Part 2: Reasons for Application
        doc.add_heading('Part 2: Reasons for Application', 1)
        
        # Add specific content based on form type
        if 'maintenance_issues' in form_type:
            doc.add_paragraph('I am applying to the Board because the landlord has not maintained the rental unit or residential complex in a good state of repair and fit for habitation.')
            
            doc.add_heading('Description of Maintenance Issues:', 2)
            doc.add_paragraph(form_data.get('issue_description', ''))
            
            doc.add_heading('Remedy Requested:', 2)
            doc.add_paragraph(form_data.get('requested_remedy', ''))
        elif 'eviction_defense' in form_type:
            doc.add_paragraph('I am applying to the Board to dispute the eviction notice served by the landlord.')
            
            doc.add_heading('Explanation of Defense:', 2)
            doc.add_paragraph(form_data.get('issue_description', ''))
            
            doc.add_heading('Remedy Requested:', 2)
            doc.add_paragraph(form_data.get('requested_remedy', ''))
        else:
            doc.add_heading('Description of Issue:', 2)
            doc.add_paragraph(form_data.get('issue_description', ''))
            
            doc.add_heading('Remedy Requested:', 2)
            doc.add_paragraph(form_data.get('requested_remedy', ''))
        
        # Part 3: Signature
        doc.add_heading('Part 3: Signature', 1)
        
        signature_table = doc.add_table(rows=2, cols=2)
        
        # Name row
        signature_table.cell(0, 0).text = 'Name:'
        signature_table.cell(0, 1).text = form_data.get('full_name', '')
        
        # Date row
        signature_table.cell(1, 0).text = 'Date:'
        today = datetime.now().strftime('%B %d, %Y')
        signature_table.cell(1, 1).text = today
        
        doc.add_paragraph()
        doc.add_paragraph('Signature: ________________________________')
        
        # Save the document
        filename = f"ltb_form_{case.id}_{uuid.uuid4().hex}.docx"
        user_id = str(case.user_id)
        upload_folder = os.path.join(os.getcwd(), 'uploads', user_id)
        os.makedirs(upload_folder, exist_ok=True)
        
        file_path = os.path.join(upload_folder, filename)
        doc.save(file_path)
        
        # Convert to PDF with watermark
        pdf_path = file_path.replace('.docx', '.pdf')
        docx_to_pdf_with_watermark(file_path, pdf_path)
        
        # Get citations
        citations = []
        
        # RTA citation
        citations.append({
            'title': 'Residential Tenancies Act, 2006',
            'citation': 'S.O. 2006, c. 17',
            'snippet': 'Section 20 requires landlords to maintain rental units in a good state of repair and fit for habitation.',
            'url': 'https://www.canlii.org/en/on/laws/stat/so-2006-c-17/latest/so-2006-c-17.html'
        })
        
        # Add relevant case law
        if 'maintenance_issues' in form_type:
            citations.append({
                'title': 'Onyskiw v. CJM Property Management Ltd.',
                'citation': '2016 ONCA 477',
                'snippet': 'The Court of Appeal for Ontario held that landlords have an ongoing obligation to maintain rental units in a good state of repair, regardless of whether the tenant was aware of issues when they moved in.',
                'url': 'https://www.canlii.org/en/on/onca/doc/2016/2016onca477/2016onca477.html'
            })
        elif 'eviction_defense' in form_type:
            citations.append({
                'title': 'Metropolitan Toronto Housing Authority v. Godwin',
                'citation': '2002 CanLII 20651 (ON LTB)',
                'snippet': 'The Landlord and Tenant Board found that a tenant is entitled to raise any issue that could be raised in a tenant application as a defense to an application for eviction.',
                'url': 'https://www.canlii.org/en/on/onltb/doc/2002/2002canlii20651/2002canlii20651.html'
            })
        
        return pdf_path, citations
    except Exception as e:
        logger.error(f"Error generating LTB form: {str(e)}")
        raise

def generate_credit_dispute(case, documents, form_type, form_data):
    """
    Generate a credit report dispute letter
    
    Args:
        case: Case model instance
        documents: List of Document model instances
        form_type: Type of form to generate
        form_data: Dictionary of form data
        
    Returns:
        tuple: (file_path, citations)
    """
    try:
        # Create a DocX document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Credit Report Dispute Letter"
        doc.core_properties.author = form_data.get('full_name', '')
        
        # Add sender information
        sender_info = f"{form_data.get('full_name', '')}\n{form_data.get('address', '')}\n{form_data.get('city', '')}, {form_data.get('province', '')} {form_data.get('postal_code', '')}\n{form_data.get('phone', '')}\n{form_data.get('email', '')}"
        doc.add_paragraph(sender_info)
        
        # Add date
        doc.add_paragraph()
        doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        doc.add_paragraph()
        
        # Add recipient information
        credit_bureau = form_data.get('credit_bureau', '')
        if credit_bureau == 'Equifax':
            recipient_info = "Equifax Canada Co.\nConsumer Relations Department\nP.O. Box 190, Station Jean-Talon\nMontreal, Quebec H1S 2Z2"
        elif credit_bureau == 'TransUnion':
            recipient_info = "TransUnion Canada\nConsumer Relations Centre\nP.O. Box 338, LCD1\nHamilton, Ontario L8L 7W2"
        else:
            recipient_info = f"{credit_bureau}\nConsumer Relations Department\nAddress Line 1\nAddress Line 2"
        
        doc.add_paragraph(recipient_info)
        doc.add_paragraph()
        
        # Add subject line
        subject = doc.add_paragraph()
        subject_run = subject.add_run(f"Subject: Credit Report Dispute - {form_data.get('full_name', '')}")
        subject_run.bold = True
        doc.add_paragraph()
        
        # Add greeting
        doc.add_paragraph("To Whom It May Concern:")
        doc.add_paragraph()
        
        # Add body
        doc.add_paragraph("I am writing to dispute information that appears on my credit report. After reviewing my credit report dated " + form_data.get('report_date', '[DATE]') + ", I have identified the following items that are inaccurate or incomplete:")
        doc.add_paragraph()
        
        # Add disputed items
        doc.add_paragraph(form_data.get('disputed_items', ''))
        doc.add_paragraph()
        
        # Add reason for dispute
        dispute_reason = form_data.get('dispute_reason', '')
        doc.add_paragraph(f"Reason for dispute: {dispute_reason}")
        
        # Add additional information if provided
        additional_info = form_data.get('additional_info', '')
        if additional_info:
            doc.add_paragraph()
            doc.add_paragraph("Additional information:")
            doc.add_paragraph(additional_info)
        
        doc.add_paragraph()
        doc.add_paragraph("As required by Section 12 of the Personal Information Protection and Electronic Documents Act (PIPEDA) and the Consumer Reporting Act, please investigate these matters and delete or correct the disputed items as soon as possible.")
        
        doc.add_paragraph()
        doc.add_paragraph("Please send me a copy of my updated credit report showing that the disputed items have been deleted or corrected. If you have any questions or need additional information, please contact me at the phone number or email address listed above.")
        
        # Add closing
        doc.add_paragraph()
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph(form_data.get('full_name', ''))
        
        # Save the document
        filename = f"credit_dispute_{case.id}_{uuid.uuid4().hex}.docx"
        user_id = str(case.user_id)
        upload_folder = os.path.join(os.getcwd(), 'uploads', user_id)
        os.makedirs(upload_folder, exist_ok=True)
        
        file_path = os.path.join(upload_folder, filename)
        doc.save(file_path)
        
        # Convert to PDF with watermark
        pdf_path = file_path.replace('.docx', '.pdf')
        docx_to_pdf_with_watermark(file_path, pdf_path)
        
        # Get citations
        citations = [{
            'title': 'Personal Information Protection and Electronic Documents Act',
            'citation': 'S.C. 2000, c. 5',
            'snippet': 'PIPEDA requires organizations to ensure that personal information they hold is accurate, complete, and up-to-date, especially when used to make decisions about an individual.',
            'url': 'https://www.canlii.org/en/ca/laws/stat/sc-2000-c-5/latest/sc-2000-c-5.html'
        }, {
            'title': 'Consumer Reporting Act',
            'citation': 'R.S.O. 1990, c. C.33',
            'snippet': 'The Act requires consumer reporting agencies to maintain procedures to ensure accuracy of information and to investigate disputed information.',
            'url': 'https://www.canlii.org/en/on/laws/stat/rso-1990-c-c33/latest/rso-1990-c-c33.html'
        }, {
            'title': 'Haskett v. Equifax Canada Inc.',
            'citation': '2003 CanLII 32896 (ON CA)',
            'snippet': 'The Court of Appeal for Ontario held that credit reporting agencies have a duty to maintain accurate information and to correct errors promptly when notified.',
            'url': 'https://www.canlii.org/en/on/onca/doc/2003/2003canlii32896/2003canlii32896.html'
        }]
        
        return pdf_path, citations
    except Exception as e:
        logger.error(f"Error generating credit dispute letter: {str(e)}")
        raise

def generate_human_rights_complaint(case, documents, form_type, form_data):
    """
    Generate a human rights complaint form
    
    Args:
        case: Case model instance
        documents: List of Document model instances
        form_type: Type of form to generate
        form_data: Dictionary of form data
        
    Returns:
        tuple: (file_path, citations)
    """
    try:
        # Create a DocX document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Human Rights Tribunal of Ontario - Application"
        doc.core_properties.author = form_data.get('full_name', '')
        
        # Add header with logo and title
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run('Human Rights Tribunal of Ontario')
        header_run.bold = True
        header_run.font.size = Pt(16)
        
        # Add title
        doc.add_heading('APPLICATION UNDER SECTION 34 OF THE HUMAN RIGHTS CODE', 0)
        doc.add_paragraph('Form 1 - Human Rights Tribunal of Ontario')
        doc.add_paragraph()
        
        # Part 1: Applicant Information
        doc.add_heading('Part 1: Applicant Information', 1)
        
        applicant_table = doc.add_table(rows=4, cols=2)
        applicant_table.style = 'Table Grid'
        
        # Name row
        applicant_table.cell(0, 0).text = 'Full Name:'
        applicant_table.cell(0, 1).text = form_data.get('full_name', '')
        
        # Address row
        applicant_table.cell(1, 0).text = 'Mailing Address:'
        applicant_table.cell(1, 1).text = f"{form_data.get('address', '')}, {form_data.get('city', '')}, {form_data.get('province', '')}, {form_data.get('postal_code', '')}"
        
        # Phone row
        applicant_table.cell(2, 0).text = 'Phone Number:'
        applicant_table.cell(2, 1).text = form_data.get('phone', 'N/A')
        
        # Email row
        applicant_table.cell(3, 0).text = 'Email:'
        applicant_table.cell(3, 1).text = form_data.get('email', 'N/A')
        
        doc.add_paragraph()
        
        # Part 2: Respondent Information
        doc.add_heading('Part 2: Respondent Information', 1)
        
        respondent_table = doc.add_table(rows=2, cols=2)
        respondent_table.style = 'Table Grid'
        
        # Name row
        respondent_table.cell(0, 0).text = 'Respondent Name:'
        respondent_table.cell(0, 1).text = form_data.get('respondent_name', '')
        
        # Address row
        respondent_table.cell(1, 0).text = 'Respondent Address:'
        respondent_table.cell(1, 1).text = form_data.get('respondent_address', '')
        
        doc.add_paragraph()
        
        # Part 3: Grounds of Discrimination
        doc.add_heading('Part 3: Grounds of Discrimination', 1)
        
        # Get discrimination grounds
        discrimination_grounds = form_data.get('discrimination_grounds', '')
        if isinstance(discrimination_grounds, str):
            discrimination_grounds = [discrimination_grounds]
        
        doc.add_paragraph('I believe that I was discriminated against based on:')
        for ground in discrimination_grounds:
            doc.add_paragraph(ground, style='List Bullet')
        
        doc.add_paragraph()
        
        # Part 4: Area of Discrimination
        doc.add_heading('Part 4: Area of Discrimination', 1)
        
        discrimination_area = form_data.get('discrimination_area', '')
        doc.add_paragraph(f"Area of Discrimination: {discrimination_area}")
        
        doc.add_paragraph()
        
        # Part 5: Description of Discrimination
        doc.add_heading('Part 5: Description of Discrimination', 1)
        
        doc.add_paragraph('Please describe what happened, when and where it happened, and who was involved. Be sure to include how you believe the ground(s) of discrimination you selected above are connected to the events you are describing:')
        doc.add_paragraph()
        
        doc.add_paragraph(form_data.get('incident_description', ''))
        
        doc.add_paragraph()
        
        # Part 6: Remedy Requested
        doc.add_heading('Part 6: Remedy Requested', 1)
        
        doc.add_paragraph('I am asking the Tribunal to:')
        doc.add_paragraph(form_data.get('requested_remedy', ''))
        
        doc.add_paragraph()
        
        # Part 7: Declaration and Signature
        doc.add_heading('Part 7: Declaration and Signature', 1)
        
        doc.add_paragraph('I declare that the information in this application is true and complete to the best of my knowledge. I understand that this application and any supporting documents will be shared with the respondent named in this application.')
        
        doc.add_paragraph()
        doc.add_paragraph('Signature: ________________________________')
        doc.add_paragraph()
        
        signature_table = doc.add_table(rows=2, cols=2)
        
        # Name row
        signature_table.cell(0, 0).text = 'Name:'
        signature_table.cell(0, 1).text = form_data.get('full_name', '')
        
        # Date row
        signature_table.cell(1, 0).text = 'Date:'
        today = datetime.now().strftime('%B %d, %Y')
        signature_table.cell(1, 1).text = today
        
        # Save the document
        filename = f"hrto_application_{case.id}_{uuid.uuid4().hex}.docx"
        user_id = str(case.user_id)
        upload_folder = os.path.join(os.getcwd(), 'uploads', user_id)
        os.makedirs(upload_folder, exist_ok=True)
        
        file_path = os.path.join(upload_folder, filename)
        doc.save(file_path)
        
        # Convert to PDF with watermark
        pdf_path = file_path.replace('.docx', '.pdf')
        docx_to_pdf_with_watermark(file_path, pdf_path)
        
        # Get citations
        citations = [{
            'title': 'Human Rights Code',
            'citation': 'R.S.O. 1990, c. H.19',
            'snippet': 'The Code prohibits discrimination in employment, housing, services, and other areas based on protected grounds such as race, age, disability, gender identity, and more.',
            'url': 'https://www.canlii.org/en/on/laws/stat/rso-1990-c-h19/latest/rso-1990-c-h19.html'
        }]
        
        # Add relevant case law based on discrimination area
        if discrimination_area == 'Employment':
            citations.append({
                'title': 'Johnstone v. Canada (Border Services Agency)',
                'citation': '2014 FCA 110',
                'snippet': 'The Federal Court of Appeal confirmed that family status includes childcare obligations and that employers have a duty to accommodate employees with these responsibilities.',
                'url': 'https://www.canlii.org/en/ca/fca/doc/2014/2014fca110/2014fca110.html'
            })
        elif discrimination_area == 'Housing':
            citations.append({
                'title': 'Québec (Commission des droits de la personne et des droits de la jeunesse) v. Montréal (City)',
                'citation': '2000 SCC 27',
                'snippet': 'The Supreme Court of Canada confirmed that "handicap" should be interpreted broadly to include both existing and perceived conditions.',
                'url': 'https://www.canlii.org/en/ca/scc/doc/2000/2000scc27/2000scc27.html'
            })
        
        return pdf_path, citations
    except Exception as e:
        logger.error(f"Error generating human rights complaint: {str(e)}")
        raise

def generate_small_claims_form(case, documents, form_type, form_data):
    """
    Generate a small claims court form
    
    Args:
        case: Case model instance
        documents: List of Document model instances
        form_type: Type of form to generate
        form_data: Dictionary of form data
        
    Returns:
        tuple: (file_path, citations)
    """
    try:
        # Create a DocX document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Small Claims Court Form"
        doc.core_properties.author = form_data.get('full_name', '')
        
        # Set up the document for a form
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # Add header with court name
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run('ONTARIO COURT OF JUSTICE\nSMALL CLAIMS COURT')
        header_run.bold = True
        header_run.font.size = Pt(16)
        
        # Add title based on form type
        if 'statement_of_claim' in form_type:
            doc.add_heading('PLAINTIFF\'S CLAIM (Form 7A)', 0)
        elif 'defense' in form_type:
            doc.add_heading('DEFENCE (Form 9A)', 0)
        else:
            doc.add_heading('SMALL CLAIMS COURT FORM', 0)
        
        # Add basic information section
        doc.add_paragraph()
        
        court_info_table = doc.add_table(rows=1, cols=2)
        court_info_table.style = 'Table Grid'
        
        court_info_table.cell(0, 0).text = 'Court File Number:'
        court_info_table.cell(0, 1).text = '(court use only)'
        
        doc.add_paragraph()
        
        # Add plaintiff/defendant information
        if 'statement_of_claim' in form_type:
            # Plaintiff section
            doc.add_heading('PLAINTIFF', 1)
            
            plaintiff_table = doc.add_table(rows=3, cols=2)
            plaintiff_table.style = 'Table Grid'
            
            plaintiff_table.cell(0, 0).text = 'Full Name:'
            plaintiff_table.cell(0, 1).text = form_data.get('full_name', '')
            
            plaintiff_table.cell(1, 0).text = 'Address:'
            plaintiff_table.cell(1, 1).text = f"{form_data.get('address', '')}, {form_data.get('city', '')}, {form_data.get('province', '')}, {form_data.get('postal_code', '')}"
            
            plaintiff_table.cell(2, 0).text = 'Phone/Email:'
            plaintiff_table.cell(2, 1).text = f"Phone: {form_data.get('phone', 'N/A')}\nEmail: {form_data.get('email', 'N/A')}"
            
            doc.add_paragraph()
            
            # Defendant section
            doc.add_heading('DEFENDANT', 1)
            
            defendant_table = doc.add_table(rows=2, cols=2)
            defendant_table.style = 'Table Grid'
            
            defendant_table.cell(0, 0).text = 'Full Name:'
            defendant_table.cell(0, 1).text = form_data.get('respondent_name', '')
            
            defendant_table.cell(1, 0).text = 'Address:'
            defendant_table.cell(1, 1).text = form_data.get('respondent_address', '')
            
            doc.add_paragraph()
            
            # Claim section
            doc.add_heading('CLAIM', 1)
            
            doc.add_paragraph('The plaintiff claims from the defendant:')
            doc.add_paragraph(form_data.get('case_description', ''))
            
            doc.add_paragraph()
            doc.add_paragraph('The plaintiff claims $_____________ for damages.')
            
            # Reasons section
            doc.add_heading('REASONS FOR CLAIM', 1)
            
            doc.add_paragraph('The following are the reasons for the claim:')
            doc.add_paragraph(form_data.get('requested_outcome', ''))
        
        elif 'defense' in form_type:
            # Defendant section
            doc.add_heading('DEFENDANT', 1)
            
            defendant_table = doc.add_table(rows=3, cols=2)
            defendant_table.style = 'Table Grid'
            
            defendant_table.cell(0, 0).text = 'Full Name:'
            defendant_table.cell(0, 1).text = form_data.get('full_name', '')
            
            defendant_table.cell(1, 0).text = 'Address:'
            defendant_table.cell(1, 1).text = f"{form_data.get('address', '')}, {form_data.get('city', '')}, {form_data.get('province', '')}, {form_data.get('postal_code', '')}"
            
            defendant_table.cell(2, 0).text = 'Phone/Email:'
            defendant_table.cell(2, 1).text = f"Phone: {form_data.get('phone', 'N/A')}\nEmail: {form_data.get('email', 'N/A')}"
            
            doc.add_paragraph()
            
            # Defence section
            doc.add_heading('DEFENCE', 1)
            
            doc.add_paragraph('I dispute the claim made against me by the plaintiff for the following reasons:')
            doc.add_paragraph(form_data.get('case_description', ''))
            
            doc.add_paragraph()
            
            # Additional details
            doc.add_heading('ADDITIONAL DETAILS', 1)
            
            doc.add_paragraph('Further information (if needed):')
            doc.add_paragraph(form_data.get('requested_outcome', ''))
        
        # Add signature section
        doc.add_heading('SIGNATURE', 1)
        
        doc.add_paragraph('Signature of ' + ('plaintiff' if 'statement_of_claim' in form_type else 'defendant') + ': ________________________________')
        doc.add_paragraph()
        
        signature_table = doc.add_table(rows=2, cols=2)
        
        signature_table.cell(0, 0).text = 'Name:'
        signature_table.cell(0, 1).text = form_data.get('full_name', '')
        
        signature_table.cell(1, 0).text = 'Date:'
        today = datetime.now().strftime('%B %d, %Y')
        signature_table.cell(1, 1).text = today
        
        # Save the document
        filename = f"small_claims_{case.id}_{uuid.uuid4().hex}.docx"
        user_id = str(case.user_id)
        upload_folder = os.path.join(os.getcwd(), 'uploads', user_id)
        os.makedirs(upload_folder, exist_ok=True)
        
        file_path = os.path.join(upload_folder, filename)
        doc.save(file_path)
        
        # Convert to PDF with watermark
        pdf_path = file_path.replace('.docx', '.pdf')
        docx_to_pdf_with_watermark(file_path, pdf_path)
        
        # Get citations
        citations = [{
            'title': 'Courts of Justice Act',
            'citation': 'R.S.O. 1990, c. C.43',
            'snippet': 'This Act establishes the Small Claims Court in Ontario and sets out its jurisdiction and procedures.',
            'url': 'https://www.canlii.org/en/on/laws/stat/rso-1990-c-c43/latest/rso-1990-c-c43.html'
        }, {
            'title': 'Small Claims Court Rules',
            'citation': 'O. Reg. 258/98',
            'snippet': 'These Rules govern the practice and procedure in the Small Claims Court, including forms, filing requirements, and trial procedures.',
            'url': 'https://www.canlii.org/en/on/laws/regu/o-reg-258-98/latest/o-reg-258-98.html'
        }]
        
        # Add relevant case law for specific claim types
        if 'statement_of_claim' in form_type:
            citations.append({
                'title': 'Garland v. Consumers\' Gas Co.',
                'citation': '2004 SCC 25',
                'snippet': 'The Supreme Court established the test for unjust enrichment as: (1) enrichment of the defendant; (2) corresponding deprivation of the plaintiff; and (3) absence of juristic reason.',
                'url': 'https://www.canlii.org/en/ca/scc/doc/2004/2004scc25/2004scc25.html'
            })
        
        return pdf_path, citations
    except Exception as e:
        logger.error(f"Error generating small claims form: {str(e)}")
        raise

def generate_child_protection_form(case, documents, form_type, form_data):
    """
    Generate a child protection form
    
    Args:
        case: Case model instance
        documents: List of Document model instances
        form_type: Type of form to generate
        form_data: Dictionary of form data
        
    Returns:
        tuple: (file_path, citations)
    """
    # Implementation for child protection forms
    # Similar structure to other forms
    try:
        # Create a generic document as placeholder
        return generate_generic_document(case, documents, form_type, form_data)
    except Exception as e:
        logger.error(f"Error generating child protection form: {str(e)}")
        raise

def generate_police_complaint(case, documents, form_type, form_data):
    """
    Generate a police complaint form
    
    Args:
        case: Case model instance
        documents: List of Document model instances
        form_type: Type of form to generate
        form_data: Dictionary of form data
        
    Returns:
        tuple: (file_path, citations)
    """
    # Implementation for police complaint forms
    # Similar structure to other forms
    try:
        # Create a generic document as placeholder
        return generate_generic_document(case, documents, form_type, form_data)
    except Exception as e:
        logger.error(f"Error generating police complaint: {str(e)}")
        raise

def generate_generic_document(case, documents, form_type, form_data):
    """
    Generate a generic legal document when specific form is not available
    
    Args:
        case: Case model instance
        documents: List of Document model instances
        form_type: Type of form to generate
        form_data: Dictionary of form data
        
    Returns:
        tuple: (file_path, citations)
    """
    try:
        # Create a DocX document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = f"Legal Document - {form_type}"
        doc.core_properties.author = form_data.get('full_name', '')
        
        # Add title
        title = form_type.replace('_', ' ').title()
        doc.add_heading(title, 0)
        
        # Add date
        doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        doc.add_paragraph()
        
        # Add applicant information
        doc.add_heading('Applicant Information', 1)
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        # Name row
        info_table.cell(0, 0).text = 'Full Name:'
        info_table.cell(0, 1).text = form_data.get('full_name', '')
        
        # Address row
        info_table.cell(1, 0).text = 'Address:'
        info_table.cell(1, 1).text = f"{form_data.get('address', '')}, {form_data.get('city', '')}, {form_data.get('province', '')}, {form_data.get('postal_code', '')}"
        
        # Phone row
        info_table.cell(2, 0).text = 'Phone:'
        info_table.cell(2, 1).text = form_data.get('phone', 'N/A')
        
        # Email row
        info_table.cell(3, 0).text = 'Email:'
        info_table.cell(3, 1).text = form_data.get('email', 'N/A')
        
        doc.add_paragraph()
        
        # Add case description
        doc.add_heading('Description', 1)
        
        doc.add_paragraph(form_data.get('case_description', ''))
        
        doc.add_paragraph()
        
        # Add requested outcome
        doc.add_heading('Requested Outcome', 1)
        
        doc.add_paragraph(form_data.get('requested_outcome', ''))
        
        doc.add_paragraph()
        
        # Add signature section
        doc.add_heading('Signature', 1)
        
        doc.add_paragraph('Signature: ________________________________')
        doc.add_paragraph()
        
        signature_table = doc.add_table(rows=2, cols=2)
        
        # Name row
        signature_table.cell(0, 0).text = 'Name:'
        signature_table.cell(0, 1).text = form_data.get('full_name', '')
        
        # Date row
        signature_table.cell(1, 0).text = 'Date:'
        today = datetime.now().strftime('%B %d, %Y')
        signature_table.cell(1, 1).text = today
        
        # Save the document
        filename = f"legal_document_{case.id}_{uuid.uuid4().hex}.docx"
        user_id = str(case.user_id)
        upload_folder = os.path.join(os.getcwd(), 'uploads', user_id)
        os.makedirs(upload_folder, exist_ok=True)
        
        file_path = os.path.join(upload_folder, filename)
        doc.save(file_path)
        
        # Convert to PDF with watermark
        pdf_path = file_path.replace('.docx', '.pdf')
        docx_to_pdf_with_watermark(file_path, pdf_path)
        
        # Generic citations
        citations = [{
            'title': 'Canadian Charter of Rights and Freedoms',
            'citation': 'Part I of the Constitution Act, 1982',
            'snippet': 'The Charter guarantees certain political rights to Canadian citizens and civil rights to everyone in Canada from the policies and actions of all areas and levels of the government.',
            'url': 'https://www.canlii.org/en/ca/laws/stat/schedule-b-to-the-canada-act-1982-uk-1982-c-11/latest/schedule-b-to-the-canada-act-1982-uk-1982-c-11.html'
        }]
        
        return pdf_path, citations
    except Exception as e:
        logger.error(f"Error generating generic document: {str(e)}")
        raise

def docx_to_pdf_with_watermark(docx_path, pdf_path):
    """
    Convert a Word document to PDF and add a watermark
    
    Args:
        docx_path (str): Path to the Word document
        pdf_path (str): Path to save the PDF
        
    Returns:
        None
    """
    try:
        # Use python-docx to convert DOCX to text (simplified conversion)
        doc = Document(docx_path)
        text_content = []
        
        for para in doc.paragraphs:
            text_content.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text_content.append(para.text)
        
        # Create a PDF with ReportLab
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = styles['Title']
        heading1_style = styles['Heading1']
        heading2_style = styles['Heading2']
        normal_style = styles['Normal']
        
        # Create watermark style
        watermark_style = ParagraphStyle(
            'Watermark',
            parent=styles['Normal'],
            fontSize=60,
            textColor=colors.lightgrey,
            alignment=1,  # center
        )
        
        # Create content flowables
        flowables = []
        
        # Add content from Word document (simplified)
        for text in text_content:
            if text.strip():
                # Try to detect headings (simplified)
                if len(text.strip()) < 50 and text.strip().isupper():
                    flowables.append(Paragraph(text, heading1_style))
                else:
                    flowables.append(Paragraph(text, normal_style))
                flowables.append(Spacer(1, 12))
        
        # Build the PDF
        doc.build(
            flowables,
            onFirstPage=lambda canvas, doc: add_watermark(canvas, doc),
            onLaterPages=lambda canvas, doc: add_watermark(canvas, doc)
        )
        
        # Save the PDF
        with open(pdf_path, 'wb') as f:
            f.write(buffer.getvalue())
    
    except Exception as e:
        logger.error(f"Error converting DOCX to PDF with watermark: {str(e)}")
        raise

def add_watermark(canvas, doc):
    """
    Add a watermark to a PDF page
    
    Args:
        canvas: ReportLab canvas
        doc: ReportLab document
    """
    canvas.saveState()
    
    # Add watermark text
    canvas.setFont('Helvetica', 60)
    canvas.setFillColorRGB(0.9, 0.9, 0.9)  # Light gray
    canvas.rotate(45)
    canvas.drawCentredString(450, -50, "SmartDispute.ai")
    canvas.drawCentredString(450, 100, "WATERMARKED")
    
    # Add footer
    canvas.setFillColorRGB(0, 0, 0)  # Black
    canvas.setFont('Helvetica', 9)
    canvas.drawString(inch, 0.5 * inch, f"Generated by SmartDispute.ai on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    canvas.drawRightString(7.5 * inch, 0.5 * inch, f"Page {doc.page}")
    
    canvas.restoreState()
def generate_referral_letter(case, documents, form_type, form_data):
    """
    Generate a referral letter
    
    Args:
        case: Case model instance
        documents: List of Document model instances
        form_type: Type of form to generate
        form_data: Dictionary of form data
        
    Returns:
        tuple: (file_path, citations)
    """
    try:
        # Create a DocX document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Referral Letter"
        doc.core_properties.author = form_data.get('full_name', '')
        
        # Add header
        doc.add_heading('REFERRAL LETTER', 0)
        doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        doc.add_paragraph()
        
        # Add recipient information
        recipient_info = f"{form_data.get('recipient_name', '')}\n{form_data.get('recipient_organization', '')}\n{form_data.get('recipient_address', '')}"
        doc.add_paragraph(recipient_info)
        doc.add_paragraph()
        
        # Add salutation
        doc.add_paragraph(f"Dear {form_data.get('recipient_name', '')}:")
        doc.add_paragraph()
        
        # Add referral content
        doc.add_paragraph(f"I am writing to refer {form_data.get('client_name', '')} for {form_data.get('service_type', '')}.")
        
        # Add case details
        doc.add_paragraph("Case Details:")
        case_details = doc.add_paragraph()
        case_details.add_run("Case Type: ").bold = True
        case_details.add_run(case.case_type.replace('-', ' ').title())
        
        # Add reason for referral
        doc.add_heading('Reason for Referral', 1)
        doc.add_paragraph(form_data.get('referral_reason', ''))
        
        # Add background information
        if form_data.get('background_info'):
            doc.add_heading('Background Information', 1)
            doc.add_paragraph(form_data.get('background_info'))
        
        # Add requested services
        doc.add_heading('Requested Services', 1)
        doc.add_paragraph(form_data.get('requested_services', ''))
        
        # Add closing
        doc.add_paragraph()
        doc.add_paragraph("Thank you for your consideration of this referral. Please contact me if you need any additional information.")
        
        # Add signature block
        doc.add_paragraph()
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph()
        doc.add_paragraph(form_data.get('full_name', ''))
        doc.add_paragraph(form_data.get('sender_title', ''))
        doc.add_paragraph(f"Phone: {form_data.get('phone', '')}")
        doc.add_paragraph(f"Email: {form_data.get('email', '')}")
        
        # Save the document
        filename = f"referral_letter_{case.id}_{uuid.uuid4().hex}.docx"
        user_id = str(case.user_id)
        upload_folder = os.path.join(os.getcwd(), 'uploads', user_id)
        os.makedirs(upload_folder, exist_ok=True)
        
        file_path = os.path.join(upload_folder, filename)
        doc.save(file_path)
        
        # Convert to PDF with watermark
        pdf_path = file_path.replace('.docx', '.pdf')
        docx_to_pdf_with_watermark(file_path, pdf_path)
        
        # Get citations
        citations = [{
            'title': 'Privacy Act',
            'citation': 'R.S.C., 1985, c. P-21',
            'snippet': 'Governs the collection, use and disclosure of personal information in the course of providing services.',
            'url': 'https://www.canlii.org/en/ca/laws/stat/rsc-1985-c-p-21/latest/rsc-1985-c-p-21.html'
        }]
        
        return pdf_path, citations
    except Exception as e:
        logger.error(f"Error generating referral letter: {str(e)}")
        raise
