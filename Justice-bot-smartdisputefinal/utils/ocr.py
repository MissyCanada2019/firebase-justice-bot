import os
import logging
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import re
import docx
from datetime import datetime

# Configure logging
logger = logging.getLogger(__name__)

def process_document(file_path, file_extension):
    """
    Process a document using OCR to extract text and metadata
    
    Args:
        file_path (str): Path to the document file
        file_extension (str): File extension (pdf, jpg, png, doc, docx)
        
    Returns:
        tuple: (extracted_text, metadata)
            - extracted_text (str): Text extracted from the document
            - metadata (dict): Metadata extracted from the document
    """
    try:
        if file_extension in ['jpg', 'jpeg', 'png']:
            return process_image(file_path)
        elif file_extension == 'pdf':
            return process_pdf(file_path)
        elif file_extension in ['doc', 'docx']:
            return process_word_document(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_extension}")
            return "", {}
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        return "", {}

def process_image(file_path):
    """
    Process an image file using Tesseract OCR
    
    Args:
        file_path (str): Path to the image file
        
    Returns:
        tuple: (extracted_text, metadata)
    """
    try:
        image = Image.open(file_path)
        
        # Extract text using Tesseract OCR
        extracted_text = pytesseract.image_to_string(image)
        
        # Extract metadata
        metadata = {
            'image_format': image.format,
            'image_size': image.size,
            'image_mode': image.mode,
            'dates': extract_dates(extracted_text),
            'names': extract_names(extracted_text),
            'addresses': extract_addresses(extracted_text),
            'phone_numbers': extract_phone_numbers(extracted_text),
            'email_addresses': extract_email_addresses(extracted_text),
        }
        
        return extracted_text, metadata
    except Exception as e:
        logger.error(f"Error processing image: {str(e)}")
        return "", {}

def process_pdf(file_path):
    """
    Process a PDF file using PyMuPDF
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        tuple: (extracted_text, metadata)
    """
    try:
        document = fitz.open(file_path)
        text = ""
        
        # Extract text from each page
        for page in document:
            text += page.get_text()
        
        # Extract metadata
        metadata = {
            'title': document.metadata.get('title', ''),
            'author': document.metadata.get('author', ''),
            'subject': document.metadata.get('subject', ''),
            'creator': document.metadata.get('creator', ''),
            'producer': document.metadata.get('producer', ''),
            'creation_date': document.metadata.get('creationDate', ''),
            'modification_date': document.metadata.get('modDate', ''),
            'page_count': len(document),
            'dates': extract_dates(text),
            'names': extract_names(text),
            'addresses': extract_addresses(text),
            'phone_numbers': extract_phone_numbers(text),
            'email_addresses': extract_email_addresses(text),
        }
        
        document.close()
        return text, metadata
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        return "", {}

def process_word_document(file_path):
    """
    Process a Word document using python-docx
    
    Args:
        file_path (str): Path to the Word document
        
    Returns:
        tuple: (extracted_text, metadata)
    """
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            'title': core_props.title if hasattr(core_props, 'title') else '',
            'author': core_props.author if hasattr(core_props, 'author') else '',
            'subject': core_props.subject if hasattr(core_props, 'subject') else '',
            'created': core_props.created.isoformat() if hasattr(core_props, 'created') and core_props.created else '',
            'modified': core_props.modified.isoformat() if hasattr(core_props, 'modified') and core_props.modified else '',
            'last_modified_by': core_props.last_modified_by if hasattr(core_props, 'last_modified_by') else '',
            'paragraph_count': len(doc.paragraphs),
            'dates': extract_dates(text),
            'names': extract_names(text),
            'addresses': extract_addresses(text),
            'phone_numbers': extract_phone_numbers(text),
            'email_addresses': extract_email_addresses(text),
        }
        
        return text, metadata
    except Exception as e:
        logger.error(f"Error processing Word document: {str(e)}")
        return "", {}

def extract_dates(text):
    """Extract dates from text"""
    # Various date formats
    date_patterns = [
        r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',  # MM/DD/YYYY or DD/MM/YYYY
        r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',  # YYYY/MM/DD
        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b',  # Month DD, YYYY
        r'\b\d{1,2} (?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}\b',  # DD Month YYYY
    ]
    
    dates = []
    for pattern in date_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        dates.extend(matches)
    
    return list(set(dates))  # Remove duplicates

def extract_names(text):
    """Extract potential names from text"""
    # Simple pattern for names (Mr./Mrs./Ms./Dr. followed by capitalized words)
    name_patterns = [
        r'\b(?:Mr|Mrs|Ms|Dr|Prof|Hon)\.\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b',
        r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2}\b'  # First Last or First Middle Last
    ]
    
    names = []
    for pattern in name_patterns:
        matches = re.findall(pattern, text)
        names.extend(matches)
    
    return list(set(names))

def extract_addresses(text):
    """Extract potential addresses from text"""
    # Simple pattern for Canadian addresses
    address_patterns = [
        r'\b\d+\s+[A-Za-z0-9\s,\.]+(?:Avenue|Ave|Boulevard|Blvd|Circle|Cir|Court|Ct|Drive|Dr|Lane|Ln|Parkway|Pkwy|Place|Pl|Plaza|Plz|Road|Rd|Square|Sq|Street|St|Way)[,\s]*(?:[A-Za-z\s]+)?(?:[,\s]*[A-Z]{2})?\s*[A-Z]\d[A-Z]\s*\d[A-Z]\d\b',
        r'\b\d+\s+[A-Za-z0-9\s,\.]+,\s*[A-Za-z\s]+,\s*[A-Z]{2},\s*[A-Z]\d[A-Z]\s*\d[A-Z]\d\b'
    ]
    
    addresses = []
    for pattern in address_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        addresses.extend(matches)
    
    return list(set(addresses))

def extract_phone_numbers(text):
    """Extract phone numbers from text"""
    # Pattern for various phone number formats
    phone_patterns = [
        r'\b\(?(?:\d{3})\)?[-.\s]?(?:\d{3})[-.\s]?(?:\d{4})\b',  # (123) 456-7890 or 123-456-7890
        r'\b\d{3}[-.\s]\d{3}[-.\s]\d{4}\b',  # 123-456-7890 or 123.456.7890
        r'\b\d{10}\b',  # 1234567890
    ]
    
    phone_numbers = []
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        phone_numbers.extend(matches)
    
    return list(set(phone_numbers))

def extract_email_addresses(text):
    """Extract email addresses from text"""
    # Pattern for email addresses
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    
    email_addresses = re.findall(email_pattern, text)
    return list(set(email_addresses))
